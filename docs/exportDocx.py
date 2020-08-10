from docxtpl import DocxTemplate
from docxtpl import InlineImage
from docx.shared import Mm
import datetime
import time
import os


curPath = os.path.abspath(os.path.dirname(__file__))
rootPath = os.path.split(curPath)[0]
# print(curPath)   c:\Users\dingp\Desktop\Gsj_Web\docs
# print(rootPath)   c:\Users\dingp\Desktop\Gsj_Web

def exportDocx():
    year="年"
    month="月"
    day="日"
    cTime=time.localtime(time.time())
    currentTime = str(cTime.tm_year)+year+str(cTime.tm_mon)+month+str(cTime.tm_mday)+day

    doc = DocxTemplate(curPath+'\mouth_report_template.docx')

    # from docx import Document
    # document = Document(curPath+'\mouth_report_template.docx')
    # run = document.tables[1].cell(0, 0).paragraphs[0].add_run()
    # run.add_picture(rootPath+'\static\\reportImg\\test.jpg')
    # run1 = document.tables[2].cell(0, 0).paragraphs[0].add_run()
    # run1.add_picture(rootPath+'\static\\reportImg\\test.jpg')
    # document.save(curPath+'\mouth_report.docx')
    t="50"

    dataDic = {
    'currentTime':currentTime,
    't1':'1,2,3,4',
    't2':"0、1、2、3、5、6、9、10、14、17、18、19、20、21、22、33、34、35、36、37、38、39、40、41、42、43、44、45、46、47、48、49、50、51、52、53、54、55、56、57、58、66、67、68、69、70、71、72、73、74号业务"+"（共" + t + "个业务）",
    'all_ipv4_mask':'11111IPv4掩码空间使用率较低，除9号、77号业务外，其余业务号使用率低于50%;1111',
    'all_ipv6_rule_mask':'222IPv6灵活空间、IPv6掩码空间的各个业务号使用率均低于50%。222',
    'ipv4_agility_content':'333IPv4灵活规则容量1000万，使用率13.44%。其中共享空间容量为654.5万，使用率为18.46%；私有空间容量345.5万，使用率3.93%，因此IPv4灵活规则总体空间充足。',
    'ipv4_agility_desc':'444已经达到申请扩容条件的业务号为：无。已经达到空间回收条件的业务号为：9、16、17、18、19、20、21、32、33、34、35、36、37、38、39、40、41、42、43、44、45、46、47、48、49、50、51、52、53、54、55、56、57、58、63、64、65、66、67、68、69、70、71、72、73、74、76、77、79、80、81、161、254、255号业务。',
    'pic_1':InlineImage(doc,rootPath+'\static\\reportImg\\1.png',width=Mm(125)),
    'pic_2':InlineImage(doc,rootPath+'\static\\reportImg\\2.png'),
    'pic_3':InlineImage(doc,rootPath+'\static\\reportImg\\3.png'),
    'pic_4':InlineImage(doc,rootPath+'\static\\reportImg\\4.png'),
    'pic_5':InlineImage(doc,rootPath+'\static\\reportImg\\5.png'),
    'pic_6':InlineImage(doc,rootPath+'\static\\reportImg\\6.png'),
    }

    doc.render(dataDic)
    doc.save(rootPath+'\static\mouth_report.docx')
exportDocx()
